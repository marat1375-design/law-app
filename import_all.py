import sys, re, sqlite3
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DB_PATH = "laws_database.db"

ARTICLE_RE    = re.compile(r'^Статья \d+[\-\d]*\.')
ARTICLE_ND_RE = re.compile(r'^Статья \d+$')   # Конституция: без точки
CHAPTER_RE    = re.compile(r'^Глава \d+[\.\s]')


def parse_by_article(path):
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    items, current_num, current_lines = [], None, []
    for text in paras:
        if ARTICLE_RE.match(text):
            if current_num:
                items.append((current_num, "\n\n".join(current_lines)))
            current_num, current_lines = text, []
        elif current_num:
            current_lines.append(text)
    if current_num:
        items.append((current_num, "\n\n".join(current_lines)))
    return items


def parse_by_article_no_dot(path):
    """Конституция: 'Статья N' без заголовка."""
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    items, current_num, current_lines = [], None, []
    for text in paras:
        if ARTICLE_ND_RE.match(text):
            if current_num:
                items.append((current_num, "\n\n".join(current_lines)))
            current_num, current_lines = text, []
        elif current_num:
            current_lines.append(text)
    if current_num:
        items.append((current_num, "\n\n".join(current_lines)))
    return items


def parse_by_chapter(path):
    """Правила: разбивка по главам."""
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    items, current_num, current_lines = [], None, []
    for text in paras:
        if CHAPTER_RE.match(text):
            if current_num:
                items.append((current_num, "\n\n".join(current_lines)))
            current_num, current_lines = text, []
        elif current_num:
            current_lines.append(text)
    if current_num:
        items.append((current_num, "\n\n".join(current_lines)))
    return items


FILES = [
    # (docx_file, law_name_in_db, delete_pattern, parser)
    (
        'k1400000235.11-06-2026.rus.docx',
        'КоАП РК (ред. 11.06.2026)',
        '%КоАП%',
        parse_by_article,
    ),
    (
        'k1400000226.rus.docx',
        'Уголовный кодекс РК от 03.07.2014 № 226-V (ред. 08.03.2026)',
        '%Уголовный кодекс%',
        parse_by_article,
    ),
    (
        'k1400000231.rus.docx',
        'Уголовно-процессуальный кодекс РК от 04.07.2014 № 231-V (ред. 08.03.2026)',
        '%процессуальный кодекс%',
        parse_by_article,
    ),
    (
        'k1500000375.rus.docx',
        'Предпринимательский кодекс РК от 29.10.2015 № 375-V (ред. 01.05.2026)',
        '%Предпринимательский кодекс%',
        parse_by_article,
    ),
    (
        'k2100000400.11-06-2026.rus.docx',
        'Экологический кодекс РК от 02.01.2021 № 400-VI (ред. 11.06.2026)',
        '%Экологический%',
        parse_by_article,
    ),
    (
        'k1700000125.11-06-2026.rus.docx',
        'Кодекс РК «О недрах и недропользовании» от 27.12.2017 № 125-VI (ред. 11.06.2026)',
        '%недрах%',
        parse_by_article,
    ),
    (
        'k2500000214.11-06-2026.rus.docx',
        'Налоговый кодекс РК от 18.07.2025 № 214-VIII (ред. 11.06.2026)',
        '%Налоговый%',
        parse_by_article,
    ),
    (
        'z1400000188.11-06-2026.rus.docx',
        'Закон РК «О гражданской защите» от 11.04.2014 № 188-V (ред. 11.06.2026)',
        '%гражданской защите%',
        parse_by_article,
    ),
    (
        'z1400000202.11-06-2026.rus.docx',
        'Закон РК «О разрешениях и уведомлениях» от 16.05.2014 № 202-V (ред. 11.06.2026)',
        '%разрешениях%',
        parse_by_article,
    ),
    (
        'z1600000442.11-06-2026.rus.docx',
        'Закон РК «Об использовании атомной энергии» от 12.01.2016 № 442-V (ред. 11.06.2026)',
        '%атомной энергии%',
        parse_by_article,
    ),
    (
        'z100000274_.rus.docx',
        'Закон РК «О защите прав потребителей» от 04.05.2010 № 274-IV (ред. 04.02.2026)',
        '%защите прав потребителей%',
        parse_by_article,
    ),
    (
        'z1500000434.rus.docx',
        'Закон РК «О государственных закупках» от 04.12.2015 № 434-V (ред. 01.01.2025)',
        '%государственных закупках%',
        parse_by_article,
    ),
    (
        'v1400010250.rus.docx',
        'Правила промышленной безопасности для объектов нефтяной и газовой отраслей (ред. 17.02.2026)',
        '%нефтяной и газовой%',
        parse_by_chapter,
    ),
    (
        'v1500011779.11-06-2026.rus.docx',
        'Правила перевозки опасных грузов автомобильным транспортом (ред. 11.06.2026)',
        '%опасных грузов%',
        parse_by_chapter,
    ),
    (
        'v2100024045.11-06-2026.rus.docx',
        'Технический регламент «Общие требования к пожарной безопасности» (ред. 11.06.2026)',
        '%пожарной безопасности%',
        parse_by_chapter,
    ),
    (
        'v2100026341.11-06-2026.rus.docx',
        'Правила управления коммунальными отходами (ред. 11.06.2026)',
        '%коммунальными отходами%',
        parse_by_chapter,
    ),
    (
        'v2300033003.11-06-2026.rus.docx',
        'Правила дорожного движения РК (ред. 11.06.2026)',
        '%дорожного движения%',
        parse_by_chapter,
    ),
    (
        'Конституция РК новая.docx',
        'Конституция Республики Казахстан (ред. 15.03.2026)',
        '%Конституция%',
        parse_by_article_no_dot,
    ),
]


def main():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    total_deleted = 0
    total_added = 0

    for docx_file, law_name, delete_pattern, parser in FILES:
        print(f"\n--- {law_name} ---")
        items = parser(docx_file)
        print(f"  Распознано: {len(items)}")

        cur.execute("DELETE FROM laws WHERE LOWER(law_name) LIKE LOWER(?)", (delete_pattern,))
        deleted = cur.rowcount
        print(f"  Удалено старых: {deleted}")

        rows = [(law_name, num, text) for num, text in items]
        cur.executemany(
            "INSERT INTO laws (law_name, article_num, text_content) VALUES (?, ?, ?)",
            rows,
        )
        print(f"  Добавлено: {len(rows)}")
        total_deleted += deleted
        total_added += len(rows)

    conn.commit()

    # Rebuild FTS5 index so search stays in sync after row deletions/insertions.
    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='laws_fts'")
    if cur.fetchone():
        cur.execute("INSERT INTO laws_fts(laws_fts) VALUES('rebuild')")
        conn.commit()
        print("FTS5 index rebuilt")

    conn.close()
    print(f"\n=== Итого: удалено {total_deleted}, добавлено {total_added} записей ===")


if __name__ == "__main__":
    main()
